try:
    from docx import Document  # Standard import
except ImportError:
    try:
        from docx.api import Document  # Fallback import
    except ImportError as ie:
        raise ImportError(
            "Failed to import 'python-docx'. Install it with: "
            "pip install python-docx"
        ) from ie

from typing import Optional
import os


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extracts text content from Word (.docx) files with robust error handling.
    Args:
        file_path: Path to the Word document
    Returns:
        Extracted text as string, or None if failed
    """
    # Validate file existence and extension
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None

    if not file_path.lower().endswith('.docx'):
        print("Invalid file format. Only .docx files are supported")
        return None

    try:
        doc = Document(file_path)
        full_text = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs
                full_text.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)

        return '\n'.join(full_text) if full_text else None

    except Exception as e:
        print(f"DOCX extraction failed: {str(e)}")
        return None


if _name_ == "_main_":
    test_file = "test.docx"  # Replace with your test file
    if os.path.exists(test_file):
        result = extract_text_from_docx(test_file)
        print("Extracted Text:" if result else "Failed to extract")
        print(result[:500] + "..." if result else "")
    else:
        print(f"Test file '{test_file}' not found")

